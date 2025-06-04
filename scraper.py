import os
import time
import json
import random
import zipfile
import pytesseract
import pandas as pd
from PIL import Image
from docx import Document
from bs4 import BeautifulSoup
from datetime import datetime
from docx.enum.text import WD_COLOR_INDEX
from playwright.sync_api import sync_playwright


CONFIG = {
    'timeout': 40000,
    'headless': False,
    'retry_attempts': 3,
    'delay_between_forms': 0,
    'viewport': {'width': 1280, 'height': 1280}
}


def crop_screenshot_to_text_bottom(raw_image_path, output_path, buffer_pixels=280):
    img = Image.open(raw_image_path)
    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    bottom_coords = [
        data['top'][i] + data['height'][i]
        for i in range(len(data['text']))
        if int(data['conf'][i]) > 30
    ]

    if not bottom_coords:
        return  # No text found; skip cropping

    max_y = min(max(bottom_coords) + buffer_pixels, img.height)
    cropped_img = img.crop((0, 0, img.width, max_y))
    cropped_img.save(output_path)
    # if os.path.exists(raw_image_path):
    #     os.remove(raw_image_path)


def setup_output_directory():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = f"Documents"
    os.makedirs(output_dir, exist_ok=True)
    return output_dir, timestamp

def clean_text(text):
    return ' '.join(text.strip().split()) if text else ""

def take_full_page_screenshot(page, filename):
    viewport_width = page.viewport_size['width']
    total_height = page.evaluate("() => document.body.scrollHeight")
    page.set_viewport_size({'width': viewport_width, 'height': total_height})
    page.locator("#documentWrap").screenshot(path=filename)

def zip_folder_and_cleanup(folder_path):
    zip_path = f"{folder_path}.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, folder_path)
                zipf.write(file_path, arcname)
    for root, dirs, files in os.walk(folder_path, topdown=False):
        for name in files:
            os.remove(os.path.join(root, name))
        for name in dirs:
            os.rmdir(os.path.join(root, name))
    os.rmdir(folder_path)
    print(f"Zipped and cleaned up: {zip_path}")

def highlight_run(run, color=WD_COLOR_INDEX.YELLOW, bold=True):
    run.bold = bold
    run.font.highlight_color = color

def scrape_document(page, url, index, output_dir):
    url_name = url['name'].replace(" ", "_")
    url_link = url['url']
    try:
        print(f"\nProcessing document {index}: {url_name}")

        for attempt in range(CONFIG['retry_attempts']):
            try:
                page.goto(url_link, timeout=CONFIG['timeout'])
                page.wait_for_selector("button.expand-button", timeout=20000)
                break
            except Exception as e:
                if attempt == CONFIG['retry_attempts'] - 1:
                    raise e
                print(f"Retry {attempt + 1} for {url_name}")
                time.sleep(2)

        form_folder = os.path.join(output_dir, url_name)
        os.makedirs(form_folder, exist_ok=True)

        form_title = page.evaluate("""() => {
            return document.querySelector('h1.document-title')?.innerText.trim() || 
                   document.title.split('|')[0].trim();
        }""")

        page.click("button.expand-button", timeout=30000, force=True)
        page.wait_for_selector("#documentWrap", timeout=40000)

        time.sleep(3)
        org_screenshot_path = os.path.join(form_folder, f"{url_name}_raw.png")
        screenshot_path = os.path.join(form_folder, f"{url_name}.png")
        take_full_page_screenshot(page, org_screenshot_path)
        crop_screenshot_to_text_bottom(org_screenshot_path, screenshot_path)

        definitions = []

        try:
            definitions_html = page.inner_html("div.definitions")
            soup = BeautifulSoup(definitions_html, 'html.parser')
            questions = soup.find_all('h3', class_='question')
            answers = soup.find_all('div', class_='answer')
            for q, a in zip(questions, answers):
                definitions.append({"question": clean_text(q.get_text()), "answer": clean_text(a.get_text())})
        except Exception as e:
            print(f"Error scraping definitions: {str(e)}")

        try:
            if page.query_selector("section#seoFaqSection"):
                faq_html = page.inner_html("section#seoFaqSection")
                faq_soup = BeautifulSoup(faq_html, 'html.parser')
                faq_items = faq_soup.select("li.faq-container")
                for item in faq_items:
                    q_div = item.find("div", class_="faq-question-container")
                    a_div = item.find("div", class_="faq-answer-container")
                    if q_div and a_div:
                        definitions.append({
                            "question": clean_text(q_div.get_text(separator=" ")),
                            "answer": clean_text(a_div.get_text(separator="\n"))
                        })
        except Exception as e:
            print(f"Error scraping FAQs: {str(e)}")

        # --- Breadcrumb Extraction ---
        breadcrumb = ""
        try:
            breadcrumb_html = page.inner_html("ol.breadcrumb-section-container")
            breadcrumb_soup = BeautifulSoup(breadcrumb_html, "html.parser")
            crumbs = breadcrumb_soup.select("li.breadcrumb span[property='name']")
            breadcrumb = " > ".join([clean_text(c.get_text()) for c in crumbs])
        except Exception as e:
            print(f"Error extracting breadcrumb: {str(e)}")

        # --- Trust Copy Extraction ---
        trust_text = ""
        try:
            trust_span = page.query_selector("span.trust-copy")
            if trust_span:
                trust_text = trust_span.inner_text().strip()
        except Exception as e:
            print(f"Error extracting trust message: {str(e)}")

        # --- Generate DOCX ---
        docx_path = os.path.join(form_folder, f"{url_name}_Writer.docx")
        doc = Document()

        # Breadcrumb
        if breadcrumb:
            p = doc.add_paragraph()
            run = p.add_run(breadcrumb)
            highlight_run(run)

        # Trust Message
        if trust_text:
            p = doc.add_paragraph()
            run = p.add_run(trust_text)
            highlight_run(run)

        doc.add_heading(form_title or url_name, 0)

        for item in definitions:
            doc.add_heading(item['question'], level=2)
            doc.add_paragraph(item['answer'])

        doc.save(docx_path)

        try:
            page.click("button.close-button", timeout=30000)
            page.wait_for_load_state("networkidle", timeout=30000)
        except:
            pass

        # zip_folder_and_cleanup(form_folder)

        return {
            "link_name": url_name,
            "form_title": form_title,
            "docx_file": docx_path.replace(form_folder, f"{form_folder}.zip"),
            "screenshot": screenshot_path.replace(form_folder, f"{form_folder}.zip")
        }

    except Exception as e:
        print(f"Failed to process {url}: {str(e)}")
        return None

def main():
    output_dir, timestamp = setup_output_directory()
    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=CONFIG['headless'],
            timeout=CONFIG['timeout'],
            args=["--disable-blink-features=AutomationControlled", "--disable-dev-shm-usage"]
        )

        context = browser.new_context(viewport=CONFIG['viewport'])
        page = context.new_page()
        
        try:
            page.goto("https://www.rocketlawyer.com/all-documents", timeout=CONFIG['timeout'])
            
            # Handle cookie consent banner
            try:
                # Wait for cookie banner to appear and click Accept
                page.wait_for_selector('#onetrust-accept-btn-handler', timeout=5000)
                page.click('#onetrust-accept-btn-handler')
                print("✅ Cookie consent accepted")
                # Give it a moment to process
                time.sleep(3)
            except:
                print("ℹ️ No cookie banner found or already accepted")    
                        
            page.wait_for_selector('div.sitemap-section', timeout=CONFIG['timeout'])
            links = page.query_selector_all('div.sitemap-section ul.sitemap-section-links li a')
            document_urls = [{"name":link.text_content(), "url": f"https://www.rocketlawyer.com{link.get_attribute('href')}", 'status': ''} for link in links]
            
            if not os.path.exists("documents.json"):
                with open("documents.json", "w", encoding="utf-8") as f:
                    json.dump(document_urls, f, indent=4)
                
            all_data = []
            for doc in document_generator("documents.json"):
                time.sleep(random.uniform(1, 2))
                
                if doc[1]['status'] != 'done':
                    document_data = scrape_document(page, doc[1], doc[0], output_dir)
                    if document_data:
                        all_data.append(document_data)
                        print(f"\n✅ Completed: {document_data['form_title']}")
                    else:
                        print(f"\n❌ Skipped : {doc[1]['url']}")
                    time.sleep(CONFIG['delay_between_forms'])

                if all_data:
                    df = pd.DataFrame([{
                        "Link Name": doc["link_name"],
                        "Form Title": doc["form_title"],
                        "ZIP File": doc["docx_file"]
                    } for doc in all_data])
                    df.to_csv(os.path.join(output_dir, "Summary.csv"), index=False)
                    print(f"\n✅ All documents scraped and zipped. Summary saved in '{output_dir}'")

        finally:
            page.close()
            context.close()
            browser.close()
            
            
def document_generator(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)

    for i, item in enumerate(data, 1):
        print(f"{i}. {item['name']} -> {item['url']}")
        yield (i, item)
        item["status"] = "done"

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)

if __name__ == "__main__":
    main()
